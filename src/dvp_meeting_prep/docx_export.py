from __future__ import annotations

from io import BytesIO
from typing import Any

import mistune
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_markdown_to_ast = mistune.create_markdown(renderer=None, plugins=["table", "strikethrough", "task_lists"])

_HEADING_STYLES = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 4",
    5: "Heading 5",
    6: "Heading 6",
}

_LIST_BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
_LIST_NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3"]


def _add_inline_children(paragraph: Any, nodes: list[dict[str, Any]], **formatting: bool) -> None:
    for node in nodes:
        _add_inline_node(paragraph, node, **formatting)


def _add_inline_node(
    paragraph: Any,
    node: dict[str, Any],
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    strike: bool = False,
    code: bool = False,
) -> None:
    node_type = node.get("type")

    if node_type == "text":
        run = paragraph.add_run(node.get("raw", ""))
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.strike = strike
        if code:
            run.font.name = "Consolas"
    elif node_type == "strong":
        _add_inline_children(paragraph, node.get("children", []), bold=True, italic=italic, underline=underline, strike=strike, code=code)
    elif node_type == "emphasis":
        _add_inline_children(paragraph, node.get("children", []), bold=bold, italic=True, underline=underline, strike=strike, code=code)
    elif node_type == "strikethrough":
        _add_inline_children(paragraph, node.get("children", []), bold=bold, italic=italic, underline=underline, strike=True, code=code)
    elif node_type == "codespan":
        run = paragraph.add_run(node.get("raw", ""))
        run.font.name = "Consolas"
        run.bold = bold
        run.italic = italic
    elif node_type in ("linebreak", "softbreak"):
        paragraph.add_run().add_break()
    elif node_type == "link":
        _add_inline_children(paragraph, node.get("children", []), bold=bold, italic=italic, underline=True, strike=strike, code=code)
    elif node_type == "image":
        alt_text = (node.get("attrs") or {}).get("alt") or "image"
        run = paragraph.add_run(f"[{alt_text}]")
        run.italic = True
    else:
        children = node.get("children")
        if children:
            _add_inline_children(paragraph, children, bold=bold, italic=italic, underline=underline, strike=strike, code=code)
        elif node.get("raw"):
            paragraph.add_run(node["raw"])


def _render_list(document: Document, list_token: dict[str, Any], depth: int = 0) -> None:
    ordered = bool((list_token.get("attrs") or {}).get("ordered"))
    styles = _LIST_NUMBER_STYLES if ordered else _LIST_BULLET_STYLES
    style = styles[min(depth, len(styles) - 1)]

    for item in list_token.get("children", []):
        prefix = ""
        if item.get("type") == "task_list_item":
            checked = bool((item.get("attrs") or {}).get("checked"))
            prefix = "☑ " if checked else "☐ "

        for child in item.get("children", []):
            child_type = child.get("type")
            if child_type in ("block_text", "paragraph"):
                paragraph = document.add_paragraph(style=style)
                if prefix:
                    paragraph.add_run(prefix)
                _add_inline_children(paragraph, child.get("children", []))
            elif child_type == "list":
                _render_list(document, child, depth=depth + 1)
            else:
                _render_block(document, child)


def _render_table(document: Document, table_token: dict[str, Any]) -> None:
    head_rows: list[list[dict[str, Any]]] = []
    body_rows: list[list[dict[str, Any]]] = []
    for section in table_token.get("children", []):
        if section.get("type") == "table_head":
            head_rows.append(section.get("children", []))
        elif section.get("type") == "table_body":
            for row in section.get("children", []):
                body_rows.append(row.get("children", []))

    if not head_rows and not body_rows:
        return

    column_count = max((len(row) for row in head_rows + body_rows), default=0)
    if column_count == 0:
        return

    table = document.add_table(rows=0, cols=column_count)
    table.style = "Light Grid Accent 1"

    for row_cells in head_rows:
        row = table.add_row()
        for index, cell in enumerate(row_cells):
            paragraph = row.cells[index].paragraphs[0]
            _add_inline_children(paragraph, cell.get("children", []), bold=True)

    for row_cells in body_rows:
        row = table.add_row()
        for index, cell in enumerate(row_cells):
            paragraph = row.cells[index].paragraphs[0]
            _add_inline_children(paragraph, cell.get("children", []))


def _render_block(document: Document, token: dict[str, Any]) -> None:
    token_type = token.get("type")

    if token_type == "heading":
        level = (token.get("attrs") or {}).get("level", 2)
        style = _HEADING_STYLES.get(level, "Heading 3")
        paragraph = document.add_paragraph(style=style)
        _add_inline_children(paragraph, token.get("children", []))
    elif token_type in ("paragraph", "block_text"):
        paragraph = document.add_paragraph()
        _add_inline_children(paragraph, token.get("children", []))
    elif token_type == "list":
        _render_list(document, token)
    elif token_type == "block_quote":
        for child in token.get("children", []):
            if child.get("type") in ("paragraph", "block_text"):
                paragraph = document.add_paragraph(style="Quote")
                _add_inline_children(paragraph, child.get("children", []))
            else:
                _render_block(document, child)
    elif token_type == "thematic_break":
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("* * *")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    elif token_type == "block_code":
        paragraph = document.add_paragraph()
        run = paragraph.add_run(token.get("raw", ""))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    elif token_type == "table":
        _render_table(document, token)
    elif token_type == "blank_line":
        return
    else:
        children = token.get("children")
        if children:
            for child in children:
                _render_block(document, child)


def markdown_to_docx_bytes(markdown_text: str, *, title: str | None = None) -> bytes:
    """Convert Markdown text into a formatted Word document, returned as bytes."""
    document = Document()

    if title:
        heading = document.add_paragraph(style="Title")
        heading.add_run(title)

    tokens = _markdown_to_ast(markdown_text) or []
    for token in tokens:
        _render_block(document, token)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
